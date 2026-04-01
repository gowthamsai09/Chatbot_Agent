from pathlib import Path
from typing import List, Dict, Set
import tempfile
import shutil
import hashlib
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import requests
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

from .settings import COLLECTION_NAME




# Document ID Utility (Phase 3)

def generate_document_id(seed: str) -> str:
    return hashlib.sha256(seed.encode("utf-8")).hexdigest()



# Disk PDF ingestion helpers
def get_indexed_pdf_sources() -> Set[str]:
    from .vector_store import get_vectorstore, get_qdrant_client
    vectorstore = get_vectorstore()
    indexed_sources = set()
    if vectorstore is None:
        return indexed_sources
    try:
        data = vectorstore.get(include=["metadatas"])
        for meta in data.get("metadatas", []):
            if meta and "source" in meta:
                indexed_sources.add(meta["source"])
    except Exception:
        pass

    return indexed_sources

# Don't need this anymore as we removed index all button from UI.
# def ingest_all_pdfs(
#     default_author: str = "Unknown",
#     default_domain: str = "general"
# ) -> Dict[str, int]:
#     indexed_sources = get_indexed_pdf_sources()
#     total_found = 0
#     newly_ingested = 0
#     skipped = 0
#     from .rag_engine import ingest_pdf

#     for pdf_path in Path(PDF_DIR).glob("*.pdf"):
#         total_found += 1
#         pdf_str = str(pdf_path)

#         if pdf_str in indexed_sources:
#             skipped += 1
#             continue

#         book_name = pdf_path.stem

#         result = ingest_pdf(
#             pdf_path=pdf_str,
#             book=book_name,
#             author=default_author,
#             domain=default_domain
#         )

#         if result.get("status") == "ingested":
#             newly_ingested += 1

#     return {
#         "pdfs_found": total_found,
#         "pdfs_ingested": newly_ingested,
#         "pdfs_skipped": skipped
#     }



# Knowledge Summary - For Qdrant
def get_knowledge_summary():
    from .vector_store import get_vectorstore, get_qdrant_client
    client = get_qdrant_client()

    documents = {}
    domains = set()

    # Scroll through all points
    offset = None
    while True:
        points, offset = client.scroll(
            collection_name=COLLECTION_NAME,
            limit=100,
            with_payload=True,
            offset=offset
        )

        for point in points:
            payload = point.payload or {}
            metadata = payload.get("metadata", {})

            doc_id = metadata.get("document_id")
            doc_name = metadata.get("document_name")
            domain = metadata.get("domain")

            if doc_id and doc_name:
                documents[doc_id] = doc_name

            if domain:
                domains.add(domain)

        if offset is None:
            break

    return {
        "documents": documents,
        "domains": sorted(domains)
    }




# TXT Upload Ingestion
def ingest_text_file(path: str, name: str, domain: str):
    from .settings import CHUNK_SIZE, CHUNK_OVERLAP, PDF_DIR
    from .vector_store import get_vectorstore, get_qdrant_client
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    chunks = splitter.split_text(text)
    docs = []

    doc_id = generate_document_id(name)

    for chunk in chunks:
        hash_value = hashlib.sha256(chunk.encode()).hexdigest()

        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "document_id": doc_id,
                    "document_name": name,
                    "domain": domain,
                    "source": "upload",
                    "chunk_id": hash_value,
                    "content_hash": hash_value
                }
            )
        )

    # vectorstore = get_vectorstore()
    # vectorstore.add_documents(docs)
    vectorstore = get_vectorstore()

    if vectorstore is None:
        return {
            "status": "error",
            "message": "Vectorstore not initialized (check HF token / env)"
        }

    vectorstore.add_documents(docs)

    return {"chunks_added": len(docs)}



# DOCX Upload Ingestion
def ingest_docx_file(path: str, name: str, domain: str):
    from .settings import CHUNK_SIZE, CHUNK_OVERLAP, PDF_DIR
    from .vector_store import get_vectorstore, get_qdrant_client
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    chunks = splitter.split_text(text)
    docs = []

    doc_id = generate_document_id(name)

    for chunk in chunks:
        hash_value = hashlib.sha256(chunk.encode()).hexdigest()

        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "document_id": doc_id,
                    "document_name": name,
                    "domain": domain,
                    "source": "upload",
                    "chunk_id": hash_value,
                    "content_hash": hash_value
                }
            )
        )

    # vectorstore = get_vectorstore()
    # vectorstore.add_documents(docs)

    vectorstore = get_vectorstore()
    if vectorstore is None:
        return {
            "status": "error",
            "message": "Vectorstore not initialized (check HF token / env)"
        }

    vectorstore.add_documents(docs)

    return {"chunks_added": len(docs)}



# Upload Dispatcher
def ingest_uploaded_document(upload_file, domain: str):
    """ Ingests a user-uploaded document (PDF / DOCX / TXT). Treats the file as ONE document to many chunks."""
    from .rag_engine import ingest_pdf
    filename = upload_file.filename.lower()

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        shutil.copyfileobj(upload_file.file, tmp)
        temp_path = tmp.name

    try:
        if filename.endswith(".pdf"):
            result = ingest_pdf(
                pdf_path=temp_path,
                book=upload_file.filename,
                author="uploaded",
                domain=domain
            )

        elif filename.endswith(".txt"):
            result = ingest_text_file(
                temp_path,
                upload_file.filename,
                domain
            )

        elif filename.endswith(".docx"):
            result = ingest_docx_file(
                temp_path,
                upload_file.filename,
                domain
            )

        else:
            return {"status": "unsupported_file"}

        return {
            "status": "ingested",
            "document": upload_file.filename,
            "details": result
        }

    finally:
        upload_file.file.close()

def ingest_url(url: str, domain: str):
    from .settings import CHUNK_SIZE, CHUNK_OVERLAP, PDF_DIR
    from .vector_store import get_vectorstore, get_qdrant_client
    """Ingests a single web page URL as ONE document to many chunks."""

    # 1. Validate URL
    if not url.startswith("http://") and not url.startswith("https://"):
        return {"status": "error", "message": "Invalid URL"}

    # 2. Fetch page
    try:
        response = requests.get(url, timeout=10, headers={
            "User-Agent": "Mozilla/5.0"
        })
    except Exception as e:
        return {"status": "error", "message": f"Request failed: {e}"}

    if response.status_code != 200:
        return {
            "status": "error",
            "message": f"Failed to fetch URL (status {response.status_code})"
        }

    # 3. Parse HTML
    soup = BeautifulSoup(response.text, "html.parser")

    # Remove scripts & styles
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    # 4. Extract title
    title = soup.title.string.strip() if soup.title and soup.title.string else url

    # 5. Extract main content
    main_content = None
    if soup.main:
        main_content = soup.main
    elif soup.article:
        main_content = soup.article
    else:
        main_content = soup.body

    if not main_content:
        return {"status": "error", "message": "No readable content found"}

    text = main_content.get_text(separator="\n", strip=True)

    if not text or len(text) < 100:
        return {"status": "error", "message": "Page content too small to index"}

    # 6. Chunk text
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    chunks = splitter.split_text(text)
    docs = []

    document_id = generate_document_id(url)

    for chunk in chunks:
        hash_value = hashlib.sha256(chunk.encode()).hexdigest()

        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "document_id": document_id,
                    "document_name": title,
                    "domain": domain,
                    "source": "url",
                    "chunk_id": hash_value,
                    "content_hash": hash_value
                }
            )
        )

    # 7. Store in vector DB
    # vectorstore = get_vectorstore()
    # vectorstore.add_documents(docs)
    vectorstore = get_vectorstore()
    if vectorstore is None:
        return {
            "status": "error",
            "message": "Vectorstore not initialized (check HF token / env)"
        }

    vectorstore.add_documents(docs)

    return {
        "status": "ingested",
        "document": title,
        "document_id": document_id,
        "chunks_added": len(docs)
    }
